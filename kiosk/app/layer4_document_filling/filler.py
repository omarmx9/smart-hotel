"""
Layer 4 — Document Filling
Responsibility: Automatically fill document templates with extracted MRZ data
Output: Filled Word document (.docx)
"""
import logging
from docx import Document
from datetime import datetime
import os

logger = logging.getLogger(__name__)


class DocumentFillingError(Exception):
    """Base exception for document filling errors"""
    def __init__(self, message, details=None):
        self.message = message
        self.details = details or {}
        super().__init__(self.message)


class TemplateNotFoundError(DocumentFillingError):
    """Template file not found"""
    def __init__(self, template_path):
        super().__init__(
            message=f"Template file not found: {template_path}",
            details={
                "template_path": template_path,
                "suggestion": "Check that the template file exists in the templates/ directory"
            }
        )


class TemplateSaveError(DocumentFillingError):
    """Failed to save filled document"""
    def __init__(self, output_path, reason):
        super().__init__(
            message=f"Failed to save document to {output_path}",
            details={
                "output_path": output_path,
                "reason": str(reason),
                "suggestion": "Check disk space and write permissions"
            }
        )


class DocumentFiller:
    """Handles automatic filling of registration card templates"""
    
    def __init__(self, template_path, output_dir="filled_documents"):
        """
        Initialize document filler
        
        Args:
            template_path: Path to the blank template .docx file
            output_dir: Directory to save filled documents
            
        Raises:
            TemplateNotFoundError: If template file doesn't exist
        """
        self.template_path = template_path
        self.output_dir = output_dir
        
        # Verify template exists
        if not os.path.exists(template_path):
            logger.error(f"Template not found: {template_path}")
            raise TemplateNotFoundError(template_path)
        
        # Create output directory if needed
        if not os.path.exists(output_dir):
            try:
                os.makedirs(output_dir)
                logger.info(f"Created output directory: {output_dir}")
            except Exception as e:
                logger.error(f"Failed to create output directory: {e}")
                raise
        
        logger.info("DocumentFiller initialized")
        logger.debug(f"  Template: {template_path}")
        logger.debug(f"  Output dir: {output_dir}")
    
    def fill_registration_card(self, mrz_data, timestamp=None):
        """
        Fill the DWA Registration Card with MRZ data
        
        Args:
            mrz_data: Dictionary containing extracted MRZ information
            timestamp: Optional timestamp for filename (uses current if None)
            
        Returns:
            dict: Contains output_path, output_filename, and timestamp
            
        Raises:
            DocumentFillingError: If document filling fails
            TemplateSaveError: If document save fails
        """
        logger.info("Starting document filling process")
        
        try:
            # Load template
            doc = Document(self.template_path)
            logger.debug("Template loaded successfully")
            
            # Extract and validate MRZ data
            surname = mrz_data.get('surname', '').strip()
            given_name = mrz_data.get('given_name', '').replace('<', ' ').strip()
            nationality = self._get_country_name(mrz_data.get('nationality_code', ''))
            passport_no = mrz_data.get('document_number', '').strip()
            birth_date = self._format_date(mrz_data.get('birth_date', ''))
            expiry_date = self._format_date(mrz_data.get('expiry_date', ''))
            issuer_country = self._get_country_name(mrz_data.get('issuer_code', ''))
            
            # Validate critical fields
            if not surname or not given_name:
                raise DocumentFillingError(
                    "Missing critical data: surname or given name",
                    details={"mrz_data": mrz_data}
                )
            
            logger.debug(f"Processing data for: {given_name} {surname}")
            
            # Define replacement mapping
            # Only fill fields available from MRZ
            # Leave other fields blank for manual entry
            replacements = {
                # MRZ fields - auto-filled
                '*Surname:': f'*Surname: {surname}',
                '*Name:': f'*Name: {given_name}',
                '*Nationality:': f'*Nationality: {nationality}',
                '*Passport No.:': f'*Passport No.: {passport_no}',
                '*Date of Birth:': f'*Date of Birth: {birth_date}',
                '*Country:': f'*Country: {issuer_country}',
                
                # Date fields
                '*From:': f'*From: {self._get_today_date()}',  # Check-in date
                '*To   :': f'*To   : {expiry_date}',  # Note: extra spaces in template
                '*Ex.  :': '*Ex.  : ..................',  # Note: extra spaces
                
                # Non-MRZ fields - leave blank for manual entry
                '*Profession:': '*Profession: ..........................................',
                '*Hometown:': '*Hometown: ..........................................',
                '*Email:': '*Email: ..........................................',
                '*Phone No.': '*Phone No. ..........................................',
                'Cabin No.': 'Cabin No. ..................',
            }
            
            # Fill the document
            self._replace_text_in_document(doc, replacements)
            
            # Generate output filename
            if timestamp is None:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            safe_name = f"{surname}_{given_name}".replace(' ', '_')[:50]
            output_filename = f"registration_card_{safe_name}_{timestamp}.docx"
            output_path = os.path.join(self.output_dir, output_filename)
            
            # Save filled document
            try:
                doc.save(output_path)
                logger.info(f"✓ Document filled and saved: {output_filename}")
            except Exception as e:
                logger.error(f"Failed to save document: {e}")
                raise TemplateSaveError(output_path, str(e))
            
            return {
                "output_path": output_path,
                "output_filename": output_filename,
                "timestamp": timestamp
            }
            
        except DocumentFillingError:
            # Re-raise our custom errors
            raise
        except Exception as e:
            logger.error(f"Error filling document: {e}")
            logger.exception("Full traceback:")
            raise DocumentFillingError(
                f"Document filling failed: {str(e)}",
                details={"error_type": type(e).__name__}
            )
    
    def _replace_text_in_document(self, doc, replacements):
        """
        Replace text in all paragraphs and tables of the document
        
        Args:
            doc: Document object
            replacements: Dictionary of {search: replace} pairs
        """
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for search_text, replace_text in replacements.items():
                if search_text in paragraph.text:
                    # Replace inline
                    for run in paragraph.runs:
                        if search_text in run.text:
                            run.text = run.text.replace(search_text, replace_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for search_text, replace_text in replacements.items():
                            if search_text in paragraph.text:
                                for run in paragraph.runs:
                                    if search_text in run.text:
                                        run.text = run.text.replace(search_text, replace_text)
    
    def _get_today_date(self):
        """
        Get today's date in DD/MM/YYYY format
        
        Returns:
            str: Today's date formatted as DD/MM/YYYY
        """
        return datetime.now().strftime('%d/%m/%Y')
    
    def _format_date(self, date_str):
        """
        Format date from YYMMDD or YYYY-MM-DD to DD/MM/YYYY
        
        Args:
            date_str: Date string in YYMMDD or YYYY-MM-DD format
            
        Returns:
            str: Formatted date DD/MM/YYYY
        """
        if not date_str:
            return ''
        
        try:
            # Try YYYY-MM-DD format first
            if '-' in date_str:
                date_obj = datetime.strptime(date_str, '%Y-%m-%d')
            # Try YYMMDD format (common in MRZ)
            elif len(date_str) == 6:
                # Assume 20xx for years 00-50, 19xx for 51-99
                year = int(date_str[0:2])
                if year <= 50:
                    full_year = 2000 + year
                else:
                    full_year = 1900 + year
                month = date_str[2:4]
                day = date_str[4:6]
                date_obj = datetime.strptime(f"{full_year}{month}{day}", '%Y%m%d')
            else:
                return date_str
            
            # Return DD/MM/YYYY
            return date_obj.strftime('%d/%m/%Y')
        except Exception as e:
            logger.warning(f"Could not format date '{date_str}': {e}")
            return date_str
    
    def _get_country_name(self, country_code):
        """
        Convert 3-letter country code to full name
        
        Args:
            country_code: ISO 3166-1 alpha-3 code (e.g., 'EGY', 'USA')
            
        Returns:
            str: Full country name
        """
        if not country_code:
            return ''
        
        # Common country codes (expanded list)
        country_map = {
            'EGY': 'Egypt',
            'USA': 'United States',
            'GBR': 'United Kingdom',
            'FRA': 'France',
            'DEU': 'Germany',
            'ITA': 'Italy',
            'ESP': 'Spain',
            'CAN': 'Canada',
            'AUS': 'Australia',
            'JPN': 'Japan',
            'CHN': 'China',
            'IND': 'India',
            'BRA': 'Brazil',
            'RUS': 'Russia',
            'SAU': 'Saudi Arabia',
            'ARE': 'United Arab Emirates',
            'TUR': 'Turkey',
            'NLD': 'Netherlands',
            'BEL': 'Belgium',
            'CHE': 'Switzerland',
            'SWE': 'Sweden',
            'NOR': 'Norway',
            'DNK': 'Denmark',
            'POL': 'Poland',
            'GRC': 'Greece',
            'PRT': 'Portugal',
            'AUT': 'Austria',
            'CZE': 'Czech Republic',
            'MEX': 'Mexico',
            'ARG': 'Argentina',
            'ZAF': 'South Africa',
            'KOR': 'South Korea',
            'SGP': 'Singapore',
            'MYS': 'Malaysia',
            'THA': 'Thailand',
            'IDN': 'Indonesia',
            'PHL': 'Philippines',
            'VNM': 'Vietnam',
            'NZL': 'New Zealand',
            'IRL': 'Ireland',
            'FIN': 'Finland',
            'ISR': 'Israel',
            'LBN': 'Lebanon',
            'JOR': 'Jordan',
            'KWT': 'Kuwait',
            'QAT': 'Qatar',
            'BHR': 'Bahrain',
            'OMN': 'Oman',
            'PAK': 'Pakistan',
            'BGD': 'Bangladesh',
            'LKA': 'Sri Lanka',
            'MAR': 'Morocco',
            'DZA': 'Algeria',
            'TUN': 'Tunisia',
            'SDN': 'Sudan',
            'YEM': 'Yemen',
            'SYR': 'Syria',
            'IRQ': 'Iraq',
            'IRN': 'Iran',
            'AFG': 'Afghanistan',
        }
        
        result = country_map.get(country_code.upper(), country_code)
        
        # If not found, return the code but log it
        if result == country_code:
            logger.debug(f"Unknown country code: {country_code}")
        
        return result