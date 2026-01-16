"""
Document Filler Module for Hotel Kiosk
Handles generating and filling the DW Registration Card (DW R.C.) with guest information.

Adapted from MRZ/app/layer4_document_filling for Django kiosk integration.
"""

import logging
import os
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)

# Check if python-docx is available
_docx_available = False
_Document = None

try:
    from docx import Document
    _Document = Document
    _docx_available = True
except ImportError:
    pass


class DocumentFillingError(Exception):
    """Base exception for document filling errors"""
    def __init__(self, message, details=None):
        self.message = message
        self.details = details or {}
        super().__init__(self.message)


class TemplateNotFoundError(DocumentFillingError):
    """Template file not found"""
    def __init__(self, template_path):
        super().__init__(
            message=f"Template file not found: {template_path}",
            details={
                "template_path": template_path,
                "suggestion": "Check that the template file exists"
            }
        )


class DocumentFiller:
    """
    Handles filling the DW Registration Card template with guest data.
    
    Can work with either .docx templates (if python-docx is available)
    or generate HTML-based documents for digital signing.
    """
    
    # Default template path (relative to project root)
    DEFAULT_TEMPLATE_PATH = os.path.join(
        os.path.dirname(os.path.dirname(__file__)),
        'MRZ', 'app', 'templates', 'DWA_Registration_Card.docx'
    )
    
    # Default output directory
    DEFAULT_OUTPUT_DIR = os.path.join(
        os.path.dirname(os.path.dirname(__file__)),
        'media', 'filled_documents'
    )
    
    def __init__(self, template_path=None, output_dir=None):
        """
        Initialize document filler.
        
        Args:
            template_path: Path to the DW R.C. template (.docx)
            output_dir: Directory to save filled documents
        """
        self.template_path = template_path or self.DEFAULT_TEMPLATE_PATH
        self.output_dir = output_dir or self.DEFAULT_OUTPUT_DIR
        
        # Create output directory if needed
        if not os.path.exists(self.output_dir):
            try:
                os.makedirs(self.output_dir)
            except Exception:
                pass
        
        # Check if template exists
        self.template_available = os.path.exists(self.template_path)
    
    def fill_registration_card(self, guest_data, timestamp=None):
        """
        Fill the DW Registration Card with guest data.
        
        Args:
            guest_data: Dictionary with guest information:
                - surname / last_name
                - name / first_name / given_name
                - nationality
                - passport_number / document_number
                - date_of_birth / birth_date
                - profession (optional)
                - hometown (optional)
                - country
                - email (optional)
                - phone (optional)
                - checkin / from_date
                - checkout / to_date
                - accompanying_guests (optional list)
                - signature_data (optional base64 image)
            timestamp: Optional timestamp for filename
            
        Returns:
            dict: Contains output_path, output_filename, timestamp, and html_preview
        """
        # Normalize field names
        data = self._normalize_guest_data(guest_data)
        
        if timestamp is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        result = {
            "timestamp": timestamp,
            "data": data,
            "html_preview": self._generate_html_preview(data),
        }
        
        # Try to fill DOCX template if available
        if _docx_available and self.template_available:
            try:
                docx_result = self._fill_docx_template(data, timestamp)
                result.update(docx_result)
            except Exception as e:
                result["docx_error"] = str(e)
        
        return result
    
    def _normalize_guest_data(self, data):
        """Normalize field names to a consistent format."""
        normalized = {
            "surname": (data.get("surname") or data.get("last_name") or "").strip(),
            "name": (data.get("name") or data.get("first_name") or data.get("given_name") or "").strip(),
            "nationality": (data.get("nationality") or "").strip(),
            "nationality_code": (data.get("nationality_code") or "").strip(),
            "passport_number": (data.get("passport_number") or data.get("document_number") or "").strip(),
            "date_of_birth": self._format_date(data.get("date_of_birth") or data.get("birth_date") or ""),
            "profession": (data.get("profession") or "").strip(),
            "hometown": (data.get("hometown") or "").strip(),
            "country": (data.get("country") or data.get("issuer_country") or "").strip(),
            "email": (data.get("email") or "").strip(),
            "phone": (data.get("phone") or "").strip(),
            "checkin": self._format_date(data.get("checkin") or data.get("from_date") or ""),
            "checkout": self._format_date(data.get("checkout") or data.get("to_date") or ""),
            "accompanying_guests": data.get("accompanying_guests") or data.get("accompany") or [],
            "signature_data": data.get("signature_data") or None,
            "signature_method": data.get("signature_method") or "physical",
        }
        
        # Set defaults for dates
        if not normalized["checkin"]:
            normalized["checkin"] = datetime.now().strftime("%Y-%m-%d")
        
        return normalized
    
    def _format_date(self, date_str):
        """Format date to display format (DD/MM/YYYY)."""
        if not date_str:
            return ""
        
        date_str = str(date_str).strip()
        
        try:
            # If already in DD/MM/YYYY format
            if "/" in date_str and len(date_str) == 10:
                return date_str
            
            # If in YYYY-MM-DD format (ISO)
            if "-" in date_str and len(date_str) == 10:
                parts = date_str.split("-")
                if len(parts) == 3:
                    return f"{parts[2]}/{parts[1]}/{parts[0]}"
            
            # YYMMDD format from MRZ
            if len(date_str) == 6 and date_str.isdigit():
                year = int(date_str[0:2])
                if year <= 50:
                    full_year = 2000 + year
                else:
                    full_year = 1900 + year
                return f"{date_str[4:6]}/{date_str[2:4]}/{full_year}"
            
            return date_str
        except Exception:
            return date_str
    
    def _fill_docx_template(self, data, timestamp):
        """Fill the DOCX template with data."""
        doc = _Document(self.template_path)
        
        # Define replacement mapping based on template placeholders
        replacements = {
            "*Surname:": f"*Surname: {data['surname']}",
            "*Name:": f"*Name: {data['name']}",
            "*Nationality:": f"*Nationality: {data['nationality']}",
            "*Passport No.:": f"*Passport No.: {data['passport_number']}",
            "*Date of Birth:": f"*Date of Birth: {data['date_of_birth']}",
            "*Country:": f"*Country: {data['country']}",
            "*Profession:": f"*Profession: {data['profession'] or '..........................................'}",
            "*Hometown:": f"*Hometown: {data['hometown'] or '..........................................'}",
            "*Email:": f"*Email: {data['email'] or '..........................................'}",
            "*Phone No.": f"*Phone No. {data['phone'] or '..........................................'}",
            "*From:": f"*From: {data['checkin']}",
            "*To   :": f"*To   : {data['checkout']}",
        }
        
        # Apply replacements
        self._replace_text_in_document(doc, replacements)
        
        # Generate output filename
        safe_name = f"{data['surname']}_{data['name']}".replace(" ", "_")[:50]
        output_filename = f"dw_registration_{safe_name}_{timestamp}.docx"
        output_path = os.path.join(self.output_dir, output_filename)
        
        # Save document
        doc.save(output_path)
        
        return {
            "output_path": output_path,
            "output_filename": output_filename,
        }
    
    def _replace_text_in_document(self, doc, replacements):
        """Replace text in document paragraphs and tables."""
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for search_text, replace_text in replacements.items():
                if search_text in paragraph.text:
                    for run in paragraph.runs:
                        if search_text in run.text:
                            run.text = run.text.replace(search_text, replace_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for search_text, replace_text in replacements.items():
                            if search_text in paragraph.text:
                                for run in paragraph.runs:
                                    if search_text in run.text:
                                        run.text = run.text.replace(search_text, replace_text)
    
    def _generate_html_preview(self, data):
        """
        Generate an HTML preview of the registration card.
        
        This is used for digital signing and on-screen preview.
        """
        accompanying_html = ""
        if data.get("accompanying_guests"):
            accompanying_html = """
            <div class="section">
                <h3>Accompanying Guests</h3>
                <table class="accompanying-table">
                    <tr><th>Name</th><th>Nationality</th><th>Passport No.</th></tr>
            """
            for i, guest in enumerate(data["accompanying_guests"], 1):
                name = guest.get("name", "")
                nationality = guest.get("nationality", "")
                passport = guest.get("passport", "")
                accompanying_html += f"<tr><td>{name}</td><td>{nationality}</td><td>{passport}</td></tr>"
            accompanying_html += "</table></div>"
        
        signature_html = ""
        if data.get("signature_data"):
            signature_html = f'<img src="{data["signature_data"]}" alt="Signature" class="signature-image" />'
        else:
            signature_html = '<div class="signature-line"></div>'
        
        html = f"""
        <div class="registration-card">
            <div class="header">
                <h1>DW Registration Card</h1>
                <p class="subtitle">Guest Registration Form</p>
            </div>
            
            <div class="section">
                <h3>Personal Information</h3>
                <div class="field-row">
                    <div class="field">
                        <label>Surname:</label>
                        <span class="value">{data['surname']}</span>
                    </div>
                    <div class="field">
                        <label>Name:</label>
                        <span class="value">{data['name']}</span>
                    </div>
                </div>
                <div class="field-row">
                    <div class="field">
                        <label>Nationality:</label>
                        <span class="value">{data['nationality']}</span>
                    </div>
                    <div class="field">
                        <label>Passport No.:</label>
                        <span class="value">{data['passport_number']}</span>
                    </div>
                </div>
                <div class="field-row">
                    <div class="field">
                        <label>Date of Birth:</label>
                        <span class="value">{data['date_of_birth']}</span>
                    </div>
                    <div class="field">
                        <label>Country:</label>
                        <span class="value">{data['country']}</span>
                    </div>
                </div>
            </div>
            
            <div class="section">
                <h3>Additional Information</h3>
                <div class="field-row">
                    <div class="field">
                        <label>Profession:</label>
                        <span class="value">{data['profession'] or '—'}</span>
                    </div>
                    <div class="field">
                        <label>Hometown:</label>
                        <span class="value">{data['hometown'] or '—'}</span>
                    </div>
                </div>
                <div class="field-row">
                    <div class="field">
                        <label>Email:</label>
                        <span class="value">{data['email'] or '—'}</span>
                    </div>
                    <div class="field">
                        <label>Phone:</label>
                        <span class="value">{data['phone'] or '—'}</span>
                    </div>
                </div>
            </div>
            
            <div class="section">
                <h3>Stay Details</h3>
                <div class="field-row">
                    <div class="field">
                        <label>Check-in:</label>
                        <span class="value">{data['checkin']}</span>
                    </div>
                    <div class="field">
                        <label>Check-out:</label>
                        <span class="value">{data['checkout'] or '—'}</span>
                    </div>
                </div>
            </div>
            
            {accompanying_html}
            
            <div class="section signature-section">
                <h3>Guest Signature</h3>
                <p class="signature-note">I confirm that all information provided is correct.</p>
                {signature_html}
                <p class="signature-date">Date: {datetime.now().strftime('%d/%m/%Y')}</p>
            </div>
        </div>
        """
        
        return html
    
    def generate_pdf(self, data, timestamp=None):
        """
        Generate a PDF version of the registration card.
        
        Returns dict with pdf_path and pdf_filename.
        """
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import mm
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError:
            return {"pdf_error": "ReportLab not installed"}
        
        if timestamp is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Output path
        safe_name = f"{data['surname']}_{data['name']}".replace(" ", "_")[:50]
        pdf_filename = f"registration_{safe_name}_{timestamp}.pdf"
        pdf_path = os.path.join(self.output_dir, pdf_filename)
        
        # Create PDF
        c = canvas.Canvas(pdf_path, pagesize=A4)
        width, height = A4
        
        # Try to use custom fonts, fallback to Helvetica
        try:
            font_regular = 'Helvetica'
            font_bold = 'Helvetica-Bold'
        except:
            font_regular = 'Helvetica'
            font_bold = 'Helvetica-Bold'
        
        # Header
        c.setFont(font_bold, 20)
        c.drawCentredString(width/2, height - 40*mm, "DW Registration Card")
        
        c.setFont(font_regular, 12)
        c.drawCentredString(width/2, height - 50*mm, "Guest Registration Form")
        
        # Line separator
        c.line(20*mm, height - 55*mm, width - 20*mm, height - 55*mm)
        
        y_pos = height - 70*mm
        line_height = 8*mm
        
        # Personal Information Section
        c.setFont(font_bold, 14)
        c.drawString(20*mm, y_pos, "Personal Information")
        y_pos -= line_height * 1.5
        
        c.setFont(font_regular, 11)
        fields = [
            ("Surname:", data['surname']),
            ("Name:", data['name']),
            ("Nationality:", data['nationality']),
            ("Passport No.:", data['passport_number']),
            ("Date of Birth:", data['date_of_birth']),
            ("Country:", data['country']),
        ]
        
        for label, value in fields:
            c.setFont(font_bold, 10)
            c.drawString(20*mm, y_pos, label)
            c.setFont(font_regular, 10)
            c.drawString(60*mm, y_pos, value or "—")
            y_pos -= line_height
        
        y_pos -= line_height * 0.5
        
        # Additional Information
        c.setFont(font_bold, 14)
        c.drawString(20*mm, y_pos, "Additional Information")
        y_pos -= line_height * 1.5
        
        c.setFont(font_regular, 11)
        extra_fields = [
            ("Profession:", data['profession'] or "—"),
            ("Hometown:", data['hometown'] or "—"),
            ("Email:", data['email'] or "—"),
            ("Phone:", data['phone'] or "—"),
        ]
        
        for label, value in extra_fields:
            c.setFont(font_bold, 10)
            c.drawString(20*mm, y_pos, label)
            c.setFont(font_regular, 10)
            c.drawString(60*mm, y_pos, value)
            y_pos -= line_height
        
        y_pos -= line_height * 0.5
        
        # Stay Details
        c.setFont(font_bold, 14)
        c.drawString(20*mm, y_pos, "Stay Details")
        y_pos -= line_height * 1.5
        
        c.setFont(font_bold, 10)
        c.drawString(20*mm, y_pos, "Check-in:")
        c.setFont(font_regular, 10)
        c.drawString(60*mm, y_pos, data['checkin'])
        
        c.setFont(font_bold, 10)
        c.drawString(100*mm, y_pos, "Check-out:")
        c.setFont(font_regular, 10)
        c.drawString(130*mm, y_pos, data['checkout'] or "—")
        y_pos -= line_height * 2
        
        # Accompanying Guests
        if data.get('accompanying_guests'):
            c.setFont(font_bold, 14)
            c.drawString(20*mm, y_pos, "Accompanying Guests")
            y_pos -= line_height * 1.5
            
            for i, guest in enumerate(data['accompanying_guests'], 1):
                c.setFont(font_regular, 10)
                guest_text = f"{i}. {guest.get('name', '')} - {guest.get('nationality', '')} - {guest.get('passport', '')}"
                c.drawString(25*mm, y_pos, guest_text)
                y_pos -= line_height
            
            y_pos -= line_height * 0.5
        
        # Signature Section
        c.setFont(font_bold, 14)
        c.drawString(20*mm, y_pos, "Guest Signature")
        y_pos -= line_height * 1.5
        
        c.setFont(font_regular, 10)
        c.drawString(20*mm, y_pos, "I confirm that all information provided is correct.")
        y_pos -= line_height * 3
        
        # Signature line
        c.line(20*mm, y_pos, 100*mm, y_pos)
        y_pos -= line_height * 0.5
        c.setFont(font_regular, 9)
        c.drawString(20*mm, y_pos, "Signature")
        
        c.drawString(120*mm, y_pos + line_height * 0.5, f"Date: {datetime.now().strftime('%d/%m/%Y')}")
        
        # Legal notice at bottom
        c.setFont(font_regular, 8)
        c.drawString(20*mm, 20*mm, "By signing this document, you agree to the hotel's terms and conditions.")
        
        c.save()
        
        return {
            "pdf_path": pdf_path,
            "pdf_filename": pdf_filename,
        }


# Singleton instance
_filler_instance = None

def get_document_filler():
    """Get the singleton document filler instance."""
    global _filler_instance
    if _filler_instance is None:
        _filler_instance = DocumentFiller()
    return _filler_instance


def fill_registration_card(guest_data, timestamp=None):
    """
    Convenience function to fill a registration card.
    
    Returns dict with html_preview, pdf_path, and optionally output_path for DOCX.
    """
    filler = get_document_filler()
    result = filler.fill_registration_card(guest_data, timestamp)
    
    # Also generate PDF
    data = filler._normalize_guest_data(guest_data)
    pdf_result = filler.generate_pdf(data, timestamp or result.get('timestamp'))
    result.update(pdf_result)
    
    return result
